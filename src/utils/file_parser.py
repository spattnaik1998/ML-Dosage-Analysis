"""
File parsing utilities for clinical documents.

Supports PDF, DOCX, TXT, and other text-based formats.
"""

import io
from pathlib import Path
from typing import Optional
import PyPDF2
from docx import Document


class FileParsingError(Exception):
    """Exception raised when file parsing fails"""
    pass


class DocumentParser:
    """Parse clinical documents from various file formats"""

    @staticmethod
    def parse_pdf(file_content: bytes) -> str:
        """
        Extract text from PDF file.

        Args:
            file_content: Raw PDF file bytes

        Returns:
            Extracted text content

        Raises:
            FileParsingError: If PDF parsing fails
        """
        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)

            text_parts = []
            for page_num, page in enumerate(pdf_reader.pages, 1):
                try:
                    text = page.extract_text()
                    if text.strip():
                        text_parts.append(f"--- Page {page_num} ---\n{text}")
                except Exception as e:
                    # Continue even if one page fails
                    text_parts.append(f"--- Page {page_num} (extraction failed) ---")

            if not text_parts:
                raise FileParsingError("No text could be extracted from PDF")

            return "\n\n".join(text_parts)

        except Exception as e:
            raise FileParsingError(f"Failed to parse PDF: {str(e)}")

    @staticmethod
    def parse_docx(file_content: bytes) -> str:
        """
        Extract text from DOCX file.

        Args:
            file_content: Raw DOCX file bytes

        Returns:
            Extracted text content

        Raises:
            FileParsingError: If DOCX parsing fails
        """
        try:
            docx_file = io.BytesIO(file_content)
            doc = Document(docx_file)

            text_parts = []

            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Extract tables
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        table_text.append(row_text)
                if table_text:
                    text_parts.append("\n[TABLE]\n" + "\n".join(table_text) + "\n[/TABLE]")

            if not text_parts:
                raise FileParsingError("No text could be extracted from DOCX")

            return "\n\n".join(text_parts)

        except Exception as e:
            raise FileParsingError(f"Failed to parse DOCX: {str(e)}")

    @staticmethod
    def parse_txt(file_content: bytes) -> str:
        """
        Extract text from plain text file.

        Args:
            file_content: Raw text file bytes

        Returns:
            Decoded text content

        Raises:
            FileParsingError: If text decoding fails
        """
        try:
            # Try common encodings
            encodings = ['utf-8', 'latin-1', 'cp1252', 'ascii']

            for encoding in encodings:
                try:
                    text = file_content.decode(encoding)
                    if text.strip():
                        return text
                except UnicodeDecodeError:
                    continue

            raise FileParsingError("Could not decode text file with any common encoding")

        except Exception as e:
            raise FileParsingError(f"Failed to parse text file: {str(e)}")

    @classmethod
    def parse_file(cls, filename: str, file_content: bytes) -> str:
        """
        Parse file based on extension.

        Args:
            filename: Original filename with extension
            file_content: Raw file bytes

        Returns:
            Extracted text content

        Raises:
            FileParsingError: If file type is unsupported or parsing fails
        """
        extension = Path(filename).suffix.lower()

        if extension == '.pdf':
            return cls.parse_pdf(file_content)
        elif extension in ['.docx', '.doc']:
            return cls.parse_docx(file_content)
        elif extension in ['.txt', '.text', '.md']:
            return cls.parse_txt(file_content)
        else:
            # Try parsing as text anyway
            try:
                return cls.parse_txt(file_content)
            except:
                raise FileParsingError(
                    f"Unsupported file type: {extension}. "
                    f"Supported formats: .pdf, .docx, .txt"
                )

    @staticmethod
    def validate_file_size(file_size: int, max_size_mb: int = 10) -> bool:
        """
        Validate file size is within limits.

        Args:
            file_size: File size in bytes
            max_size_mb: Maximum allowed size in MB

        Returns:
            True if valid, False otherwise
        """
        max_bytes = max_size_mb * 1024 * 1024
        return file_size <= max_bytes

    @staticmethod
    def estimate_case_studies(text: str) -> int:
        """
        Estimate number of case studies in document.

        Args:
            text: Document text

        Returns:
            Estimated number of case studies
        """
        # Look for common case study indicators
        indicators = [
            'patient is a',
            'case study',
            'clinical case',
            'patient #',
            'patient id',
            'case #',
            'case:',
            'admission:',
        ]

        count = 0
        text_lower = text.lower()

        for indicator in indicators:
            count = max(count, text_lower.count(indicator))

        # Default to at least 1 if text exists
        return max(1, count)
